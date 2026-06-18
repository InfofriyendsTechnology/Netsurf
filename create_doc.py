import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Title
title = doc.add_heading('Business Overview: Kishorbhai Babariya', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# ==========================================
# PART 1: ENGLISH
# ==========================================
doc.add_heading('PART 1: ENGLISH', level=1)

doc.add_heading('1. What is our business and what is the purpose of our business?', level=2)
p = doc.add_paragraph()
p.add_run("Your business operates within the ")
p.add_run("Direct Selling (Multi-Level Marketing)").bold = True
p.add_run(" industry. The core purpose of your business is to promote a healthy, natural, and sustainable lifestyle. You do this by providing high-quality, herbal, and organic products that solve real-life problems—ranging from personal health issues to agricultural challenges. Additionally, the business model empowers individuals by giving them an opportunity to become independent distributors, build a team, and create a sustainable source of income.")

doc.add_heading('2. What we are doing, which is our company?', level=2)
p = doc.add_paragraph()
p.add_run("Your company is ")
p.add_run("Netsurf Network").bold = True
p.add_run(" (part of Netsurf Communications Pvt. Ltd.), one of the leading direct-selling companies in India.\n\nThrough your YouTube channel (@KishorbhaiBabariya), you are actively working as an independent distributor and promoter for the company. You are educating people by sharing real-life testimonials and results of using Netsurf products—particularly focusing on Ayurvedic health treatments, nutritional improvements, and organic farming techniques (like increasing crop yield and quality). You are also likely recruiting and training a network of individuals to expand the business.")

doc.add_heading('3. What is our product and who am I?', level=2)
p = doc.add_paragraph()
p.add_run("Your Products:").bold = True
doc.add_paragraph("Netsurf offers a diverse range of natural, herbal, and result-oriented products, primarily divided into five categories:")
doc.add_paragraph("Health & Wellness (Naturamore): Nutritional supplements for overall well-being and managing lifestyle-related health issues.", style='List Bullet')
doc.add_paragraph("Agriculture (Biofit): Organic farming solutions, bio-pesticides, and crop-care products to improve farming yield and soil health safely.", style='List Bullet')
doc.add_paragraph("Personal Care (Herbs & More): Herbal skin, hair, and body care products.", style='List Bullet')
doc.add_paragraph("Home Care (Clean & More): Eco-friendly, enzyme-based cleaning solutions.", style='List Bullet')
doc.add_paragraph("Colour Cosmetics (Rang Dé): Makeup products formulated with safe ingredients.", style='List Bullet')

p = doc.add_paragraph()
p.add_run("Who You Are:").bold = True
p = doc.add_paragraph()
p.add_run("You are ")
p.add_run("Kishorbhai Babariya").bold = True
p.add_run(", an independent business owner, distributor, and leader affiliated with Netsurf Network. Through your digital presence, you act as a consultant and guide, helping farmers transition to organic agriculture and helping individuals find natural health solutions through Ayurveda and Netsurf's product lines.")

doc.add_page_break()

# ==========================================
# PART 2: GUJARATI
# ==========================================
doc.add_heading('PART 2: GUJARATI (ગુજરાતી)', level=1)

doc.add_heading('૧. આપણો વ્યવસાય શું છે અને તેનો હેતુ શું છે?', level=2)
p = doc.add_paragraph("તમારો વ્યવસાય ")
p.add_run("ડાયરેક્ટ સેલિંગ (નેટવર્ક માર્કેટિંગ)").bold = True
p.add_run(" ક્ષેત્રમાં છે. આ વ્યવસાયનો મુખ્ય હેતુ લોકોને સ્વસ્થ, કુદરતી અને સારું જીવન જીવવામાં મદદ કરવાનો છે. તમે ઉચ્ચ ગુણવત્તાવાળી, હર્બલ અને ઓર્ગેનિક પ્રોડક્ટ્સ પૂરી પાડીને લોકોના સ્વાસ્થ્ય અને ખેતીને લગતી સમસ્યાઓનું સચોટ સમાધાન કરો છો. આ સાથે જ, તમે અન્ય લોકોને એક સ્વતંત્ર ડિસ્ટ્રીબ્યુટર (વિક્રેતા) બનવાની શાનદાર તક આપીને તેમને આર્થિક રીતે સદ્ધર, આત્મનિર્ભર અને રોજગારી મેળવવામાં મદદ કરવાનો હેતુ ધરાવો છો.")

doc.add_heading('૨. આપણે શું કરી રહ્યા છીએ, અને આપણી કંપની કઈ છે?', level=2)
p = doc.add_paragraph("તમારી કંપનીનું નામ ")
p.add_run("Netsurf Network").bold = True
p.add_run(" (નેટસર્ફ કોમ્યુનિકેશન્સ પ્રાઈવેટ લિમિટેડ) છે, જે ભારતની અગ્રણી ડાયરેક્ટ સેલિંગ કંપનીઓમાંની એક છે.\n\nતમારી યુટ્યુબ ચેનલના માધ્યમથી, તમે આ કંપનીના એક સ્વતંત્ર ડિસ્ટ્રીબ્યુટર અને પ્રમોટર તરીકે કામ કરી રહ્યા છો. તમે આયુર્વેદિક સ્વાસ્થ્ય સારવાર, ન્યુટ્રિશનલ સપ્લિમેન્ટ્સ અને ઓર્ગેનિક ખેતી (જેમ કે જમીન સુધારવી અને પાકનું ઉત્પાદન વધારવું) માટે નેટસર્ફની પ્રોડક્ટ્સના સચોટ પરિણામો અને લોકોના અનુભવો (ટેસ્ટિમોનિયલ્સ) ના વિડીયો શેર કરીને ખેડૂતો અને સામાન્ય લોકોને માહિતગાર કરી રહ્યા છો. આ ઉપરાંત, તમે તમારી ટીમ અને નેટવર્ક વધારવા માટે લોકોને તાલીમ પણ આપો છો.")

doc.add_heading('૩. આપણી પ્રોડક્ટ્સ કઈ છે અને હું કોણ છું?', level=2)
p = doc.add_paragraph()
p.add_run("તમારી પ્રોડક્ટ્સ:").bold = True
doc.add_paragraph("નેટસર્ફ મુખ્યત્વે પાંચ કેટેગરીમાં ૧૦૦% કુદરતી, હર્બલ અને ઉત્તમ પરિણામ આપતી પ્રોડક્ટ્સ બનાવે છે:")
doc.add_paragraph("સ્વાસ્થ્ય અને સુખાકારી (Naturamore): શારીરિક સ્વાસ્થ્ય સુધારવા, રોગપ્રતિકારક શક્તિ વધારવા અને બીમારીઓથી બચવા માટેના ન્યુટ્રિશનલ સપ્લિમેન્ટ્સ.", style='List Bullet')
doc.add_paragraph("ખેતી (Biofit): ઓર્ગેનિક ખેતી માટેની પ્રોડક્ટ્સ, જે રાસાયણિક ખાતરનો ખર્ચ ઘટાડે છે, જમીન સુધારે છે અને પાકનું ઉત્પાદન વધારે છે.", style='List Bullet')
doc.add_paragraph("પર્સનલ કેર (Herbs & More): ત્વચા, વાળ અને શરીરની સંભાળ માટેની ઉત્તમ હર્બલ પ્રોડક્ટ્સ.", style='List Bullet')
doc.add_paragraph("હોમ કેર (Clean & More): ઘરની સાફ-સફાઈ માટેની ઈકો-ફ્રેન્ડલી અને સુરક્ષિત પ્રોડક્ટ્સ.", style='List Bullet')
doc.add_paragraph("કલર કોસ્મેટિક્સ (Rang Dé): હાનિકારક કેમિકલ વગરના સુરક્ષિત તત્વોમાંથી બનાવેલ મેકઅપ પ્રોડક્ટ્સ.", style='List Bullet')

p = doc.add_paragraph()
p.add_run("તમે કોણ છો:").bold = True
p = doc.add_paragraph("તમે ")
p.add_run("કિશોરભાઈ બાબરિયા").bold = True
p.add_run(" છો. તમે નેટસર્ફ નેટવર્ક સાથે જોડાયેલા એક સફળ, સ્વતંત્ર બિઝનેસ ઓનર, ડિસ્ટ્રીબ્યુટર અને લીડર છો. તમારા ડિજિટલ માધ્યમ (YouTube ચેનલ) દ્વારા તમે એક સાચા સલાહકાર અને માર્ગદર્શક તરીકે કામ કરી રહ્યા છો, જે ખેડૂતોને રાસાયણિક ખેતી છોડીને ઓર્ગેનિક ખેતી તરફ વળવામાં અને સામાન્ય લોકોને આયુર્વેદિક તેમજ નેટસર્ફની પ્રોડક્ટ્સ દ્વારા તેમના સ્વાસ્થ્યની જાળવણી કરવામાં મદદ કરે છે.")

doc.add_page_break()

# ==========================================
# PART 3: HINDI
# ==========================================
doc.add_heading('PART 3: HINDI (हिंदी)', level=1)

doc.add_heading('1. हमारा व्यवसाय क्या है और हमारे व्यवसाय का उद्देश्य क्या है?', level=2)
p = doc.add_paragraph("आपका व्यवसाय ")
p.add_run("डायरेक्ट सेलिंग (नेटवर्क मार्केटिंग)").bold = True
p.add_run(" उद्योग से जुड़ा है। आपके व्यवसाय का मुख्य उद्देश्य एक स्वस्थ, प्राकृतिक और टिकाऊ जीवन शैली को बढ़ावा देना है। आप उच्च गुणवत्ता वाले, हर्बल और जैविक (ऑर्गेनिक) उत्पाद प्रदान करके लोगों के स्वास्थ्य और कृषि से जुड़ी समस्याओं का समाधान करते हैं। इसके साथ ही, यह व्यवसाय मॉडल लोगों को एक स्वतंत्र वितरक (डिस्ट्रीब्यूटर) बनने का अवसर देकर उन्हें आर्थिक रूप से सशक्त और आत्मनिर्भर बनाने में मदद करता है।")

doc.add_heading('2. हम क्या कर रहे हैं, और हमारी कंपनी कौन सी है?', level=2)
p = doc.add_paragraph("आपकी कंपनी ")
p.add_run("Netsurf Network").bold = True
p.add_run(" (नेट्सर्फ कम्युनिकेशंस प्राइवेट लिमिटेड) है, जो भारत की अग्रणी डायरेक्ट सेलिंग कंपनियों में से एक है।\n\nअपने YouTube चैनल के माध्यम से, आप कंपनी के लिए एक स्वतंत्र वितरक और प्रमोटर के रूप में सक्रिय रूप से काम कर रहे हैं। आप आयुर्वेदिक स्वास्थ्य उपचार, पोषण की खुराक (न्यूट्रीशनल सप्लीमेंट्स) और जैविक खेती की तकनीकों (जैसे फसल की पैदावार और गुणवत्ता बढ़ाना) के लिए नेट्सर्फ उत्पादों के उपयोग के वास्तविक परिणाम और लोगों के अनुभव साझा करके किसानों और आम जनता को जागरूक कर रहे हैं। आप व्यवसाय के विस्तार के लिए लोगों को अपनी टीम में जोड़कर उन्हें प्रशिक्षित भी कर रहे हैं।")

doc.add_heading('3. हमारे उत्पाद क्या हैं और मैं कौन हूँ?', level=2)
p = doc.add_paragraph()
p.add_run("आपके उत्पाद:").bold = True
doc.add_paragraph("नेट्सर्फ मुख्य रूप से पांच श्रेणियों में प्राकृतिक, हर्बल और परिणाम-उन्मुख उत्पाद प्रदान करता है:")
doc.add_paragraph("स्वास्थ्य और कल्याण (Naturamore): समग्र स्वास्थ्य के लिए पोषण की खुराक (न्यूट्रीशनल सप्लीमेंट्स) और जीवनशैली से जुड़ी बीमारियों से बचाव के लिए उत्पाद।", style='List Bullet')
doc.add_paragraph("कृषि (Biofit): जैविक खेती के समाधान, जो रसायनों का खर्च कम करते हैं, मिट्टी की सेहत सुधारते हैं और फसल की पैदावार को सुरक्षित रूप से बढ़ाते हैं।", style='List Bullet')
doc.add_paragraph("पर्सनल केयर (Herbs & More): त्वचा, बालों और शरीर की देखभाल के लिए हर्बल उत्पाद।", style='List Bullet')
doc.add_paragraph("होम केयर (Clean & More): घर की साफ-सफाई के लिए पर्यावरण के अनुकूल (इको-फ्रेंडली) उत्पाद।", style='List Bullet')
doc.add_paragraph("कलर कॉस्मेटिक्स (Rang Dé): सुरक्षित तत्वों से बने मेकअप उत्पाद।", style='List Bullet')

p = doc.add_paragraph()
p.add_run("आप कौन हैं:").bold = True
p = doc.add_paragraph("आप ")
p.add_run("किशोरभाई बाबरिया").bold = True
p.add_run(" हैं, जो नेट्सर्फ नेटवर्क से जुड़े एक स्वतंत्र बिजनेस ओनर, वितरक (डिस्ट्रीब्यूटर) और लीडर हैं। अपनी डिजिटल उपस्थिति (YouTube) के माध्यम से, आप एक सलाहकार और मार्गदर्शक के रूप में कार्य करते हैं, जो किसानों को जैविक कृषि की ओर बढ़ने में मदद करते हैं और आयुर्वेद व नेट्सर्फ उत्पादों के माध्यम से लोगों को उनके स्वास्थ्य के लिए प्राकृतिक समाधान खोजने में मार्गदर्शन करते हैं।")

doc.save(r'c:\Users\SIS\OneDrive\Desktop\IMPORTANT\Netsurf\Kishorbhai_Babariya_Business_Details_V2.docx')
